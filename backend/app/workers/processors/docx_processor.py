import asyncio
import logging
from typing import Dict, Any
from app.workers.processors.base_processor import BaseProcessor

logger = logging.getLogger(__name__)


class DOCXProcessor(BaseProcessor):

    async def parse(self, file_path: str) -> Dict[str, Any]:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._parse_sync, file_path)

    def _parse_sync(self, file_path: str) -> Dict[str, Any]:
        try:
            import docx
            doc = docx.Document(file_path)

            # Extract paragraphs + tables
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)

            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            }
            return {"text": "\n".join(parts), "metadata": metadata}
        except Exception as e:
            logger.error(f"DOCX parse error: {e}")
            raise

    async def extract_structured_data(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        text = parsed_data["text"]
        return {
            "text": text,
            "title": self.extract_title(text),
            "category": self.detect_category(text),
            "summary": self.extract_summary(text),
            "keywords": self.extract_keywords(text),
            "metadata": parsed_data["metadata"],
        }
